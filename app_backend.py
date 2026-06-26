import atexit
import json
import os
import subprocess
import sys
import tempfile
import threading
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import cv2
import easyocr
import httpx
import numpy as np
import pypdfium2 as pdfium
from dotenv import load_dotenv
from openai import OpenAI

from llmops.local_extraction import extract_invoice_fields_local, parse_model_json_or_fallback
from llmops.registry import load_prompt_registry
from llmops.schema import validate_invoice_fields
from llmops.telemetry import calculate_usage_cost, get_default_telemetry, normalize_usage
from llmops.tracing import write_trace_record


class Milestone1NotebookAPI:
    """Backend API adapted from milestone invoice OCR and extraction flows."""

    SCALAR_FIELDS = [
        "invoice_number",
        "invoice_date",
        "due_date",
        "po_number",
        "payment_terms",
        "vendor_name",
        "vendor_tax_id",
        "customer_name",
        "customer_tax_id",
        "subtotal",
        "tax",
        "total",
        "currency",
    ]
    ORDER_ITEM_FIELDS = [
        "line_no",
        "description",
        "qty",
        "unit",
        "unit_price",
        "net_amount",
        "tax_rate",
        "gross_amount",
    ]

    def __init__(self, workspace_root: Path):
        self.workspace_root = workspace_root
        self.dataset_root = workspace_root / "Datasets"
        self.output_dir = workspace_root / "outputs"
        self.output_dir.mkdir(exist_ok=True)
        self.trace_path = self.output_dir / "llmops_traces.jsonl"

        self.easyocr_model_dir = self._resolve_env_path("EASYOCR_MODEL_DIR", workspace_root)
        self.craft_model_path = self.easyocr_model_dir / "craft_mlt_25k.pth"
        self.english_model_path = self.easyocr_model_dir / "english_g2.pth"
        self.ocr_available = self.craft_model_path.exists() and self.english_model_path.exists()
        self.ocr_unavailable_reason: Optional[str] = None
        self.reader = None
        if not self.ocr_available:
            self.ocr_unavailable_reason = (
                "Required EasyOCR model files not found: craft_mlt_25k.pth and english_g2.pth"
            )
        else:
            self.reader = easyocr.Reader(
                ["en"],
                gpu=False,
                model_storage_directory=str(self.easyocr_model_dir),
                detect_network="craft",
                recog_network="english_g2",
                download_enabled=False,
                verbose=False,
            )

        env_file = workspace_root / ".env"
        load_dotenv(env_file)
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_api_base = os.getenv(
            "OPENAI_API_BASE", "https://aibe.mygreatlearning.com/openai/v1"
        )
        self.model_name = "gpt-4o-mini"
        self.field_extractor_mode = os.getenv("FIELD_EXTRACTOR_MODE", "auto").strip().lower()
        if self.field_extractor_mode not in {"auto", "gpt", "qwen"}:
            self.field_extractor_mode = "auto"

        self.prompt_version = os.getenv("LLMOPS_PROMPT_VERSION", "v2").strip() or "v2"
        self.schema_version = os.getenv("LLMOPS_SCHEMA_VERSION", "v2").strip() or "v2"
        self.trace_include_text = os.getenv("LLMOPS_TRACE_TEXT", "false").strip().lower() == "true"

        self.prompt_registry = load_prompt_registry(workspace_root)
        self.invoice_prompt_entry = self.prompt_registry.get_invoice_entry(self.prompt_version)
        self.invoice_prompt = self.invoice_prompt_entry.prompt_path.read_text(encoding="utf-8")
        self.invoice_schema_path = self.invoice_prompt_entry.schema_path
        self.pricing_path = self._resolve_env_path(
            "LLMOPS_PRICING_FILE",
            workspace_root / "configs" / "model_pricing.yaml",
        )
        self.telemetry = get_default_telemetry(self.pricing_path)

        self.qwen_model_dir = self._resolve_env_path(
            "QWEN_MODEL_DIR",
            workspace_root / "qwen3-vl-8b-instruct",
        )
        self.qwen_model_name = "Qwen3-VL-8B-Instruct (local)"
        self._qwen_model = None
        self._qwen_processor = None
        self._qwen_device = "cpu"
        self._qwen_load_error: Optional[str] = None

        self.layout_worker_python = self._resolve_layout_worker_python()
        self.layout_worker_script = workspace_root / "scripts" / "layout_worker.py"
        self._layout_proc = None
        self._layout_lock = threading.Lock()
        self._layout_error: Optional[str] = None

        self.openai_client = None
        if self.openai_api_key:
            http_client = httpx.Client(verify=False)
            self.openai_client = OpenAI(
                api_key=self.openai_api_key,
                base_url=self.openai_api_base,
                http_client=http_client,
            )

        metrics_port = (os.getenv("PROMETHEUS_METRICS_PORT") or "").strip()
        if metrics_port:
            self.telemetry.start_server(int(metrics_port))

    def _resolve_env_path(self, env_name: str, default_path: Path) -> Path:
        configured = (os.getenv(env_name) or "").strip()
        if not configured:
            return default_path
        candidate = Path(configured)
        return candidate if candidate.is_absolute() else self.workspace_root / candidate

    def _resolve_layout_worker_python(self) -> Path:
        explicit = os.getenv("LAYOUT_WORKER_PYTHON")
        if explicit:
            return Path(explicit)

        candidates = [
            self.workspace_root / ".venv-layout" / "Scripts" / "python.exe",
            self.workspace_root / ".venv-layout" / "bin" / "python",
            Path(sys.executable),
        ]
        for candidate in candidates:
            if candidate.exists():
                return candidate
        return candidates[0]

    def _ensure_layout_worker(self) -> bool:
        if self._layout_proc is not None and self._layout_proc.poll() is None:
            return True
        if self._layout_error is not None:
            return False
        if not self.layout_worker_script.exists():
            self._layout_error = f"layout_worker_missing: worker={self.layout_worker_script.exists()}"
            return False
        if not self.layout_worker_python.exists():
            self._layout_error = f"layout_worker_missing: python={self.layout_worker_python}"
            return False
        try:
            self._layout_proc = subprocess.Popen(
                [str(self.layout_worker_python), str(self.layout_worker_script), "--batch"],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.DEVNULL,
                text=True,
                bufsize=1,
            )
            atexit.register(self._close_layout_worker)
            return True
        except Exception as exc:
            self._layout_error = f"layout_worker_start_failed: {exc}"
            return False

    def __del__(self):
        self._close_layout_worker()

    def _close_layout_worker(self) -> None:
        try:
            if self._layout_proc is not None and self._layout_proc.poll() is None:
                if self._layout_proc.stdin:
                    self._layout_proc.stdin.close()
                self._layout_proc.terminate()
                self._layout_proc.wait(timeout=5)
        except Exception:
            pass

    def detect_layout_regions(self, image_rgb: np.ndarray) -> Dict[str, Any]:
        if not self._ensure_layout_worker():
            return {
                "status": "layout_worker_unavailable",
                "regions": [],
                "note": self._layout_error,
            }

        fd, tmp_path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        try:
            bgr = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2BGR)
            cv2.imwrite(tmp_path, bgr)
            with self._layout_lock:
                assert self._layout_proc is not None
                assert self._layout_proc.stdin is not None
                assert self._layout_proc.stdout is not None
                self._layout_proc.stdin.write(tmp_path + "\n")
                self._layout_proc.stdin.flush()
                line = self._layout_proc.stdout.readline()
            if not line:
                return {"status": "layout_worker_dead", "regions": []}
            return json.loads(line)
        except Exception as exc:
            return {"status": "layout_worker_error", "regions": [], "note": str(exc)}
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    @staticmethod
    def _sort_regions_reading_order(regions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        def _key(region: Dict[str, Any]) -> Tuple[int, int]:
            x1, y1, _, _ = region.get("bbox", [0, 0, 0, 0])
            return int(y1), int(x1)

        return sorted(regions, key=_key)

    def quality_metrics(self, image_rgb: np.ndarray) -> Dict[str, float]:
        gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
        blur = float(cv2.Laplacian(gray, cv2.CV_64F).var())
        brightness = float(np.mean(gray))
        contrast = float(np.std(gray))
        return {"blur_var": blur, "brightness": brightness, "contrast": contrast}

    def preprocess_for_ocr(self, image_rgb: np.ndarray) -> np.ndarray:
        gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        thresholded = cv2.adaptiveThreshold(
            denoised,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            5,
        )
        return cv2.cvtColor(thresholded, cv2.COLOR_GRAY2RGB)

    def easyocr_on_image_array(self, image_rgb: np.ndarray) -> Tuple[str, float, int]:
        if self.reader is None:
            raise RuntimeError(self.ocr_unavailable_reason or "EasyOCR reader unavailable")
        results = self.reader.readtext(image_rgb)
        texts, confidences = [], []
        for item in results:
            if len(item) < 3:
                continue
            text = str(item[1]).strip()
            if not text:
                continue
            texts.append(text)
            confidences.append(float(item[2]))

        merged_text = "\n".join(texts)
        avg_confidence = float(sum(confidences) / len(confidences)) if confidences else 0.0
        return merged_text, avg_confidence, len(texts)

    def ocr_image_layout_aware(self, image_rgb: np.ndarray) -> Tuple[str, float, int, Dict[str, Any]]:
        preprocessed = self.preprocess_for_ocr(image_rgb)
        full_text, full_conf, full_det = self.easyocr_on_image_array(preprocessed)

        layout = self.detect_layout_regions(image_rgb)
        regions = layout.get("regions", []) if isinstance(layout, dict) else []
        if not regions:
            return full_text, full_conf, full_det, {
                "layout_status": layout.get("status", "unknown") if isinstance(layout, dict) else "unknown",
                "layout_regions": 0,
                "layout_crop_chunks": 0,
            }

        height, width = image_rgb.shape[:2]
        crop_texts: List[str] = []
        crop_confs: List[float] = []
        crop_dets = 0

        ordered = self._sort_regions_reading_order(regions)
        for index, region in enumerate(ordered[:30], start=1):
            bbox = region.get("bbox", [])
            if len(bbox) != 4:
                continue
            x1, y1, x2, y2 = [int(value) for value in bbox]
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(width, x2), min(height, y2)
            if x2 - x1 < 12 or y2 - y1 < 12:
                continue

            crop = image_rgb[y1:y2, x1:x2]
            crop_preprocessed = self.preprocess_for_ocr(crop)
            crop_text, crop_conf, crop_det = self.easyocr_on_image_array(crop_preprocessed)
            if crop_det <= 0 or not crop_text.strip():
                continue

            crop_texts.append(f"[REGION {index}]\n{crop_text}")
            crop_confs.append(crop_conf)
            crop_dets += crop_det

        if not crop_texts:
            return full_text, full_conf, full_det, {
                "layout_status": layout.get("status", "unknown"),
                "layout_regions": len(regions),
                "layout_crop_chunks": 0,
            }

        merged_text = (
            "=== FULL PAGE OCR ===\n"
            + full_text
            + "\n\n=== LAYOUT REGION OCR ===\n"
            + "\n\n".join(crop_texts)
        )
        conf_values = [value for value in [full_conf, *crop_confs] if isinstance(value, (int, float))]
        merged_conf = round(float(sum(conf_values) / len(conf_values)), 2) if conf_values else full_conf
        merged_det = int(full_det + crop_dets)
        return merged_text, merged_conf, merged_det, {
            "layout_status": layout.get("status", "unknown"),
            "layout_regions": len(regions),
            "layout_crop_chunks": len(crop_texts),
        }

    def extract_fields_local(self, text: str) -> Dict[str, Any]:
        return extract_invoice_fields_local(text)

    def _local_with_reason(
        self,
        text: str,
        reason: str,
        detail: Optional[str] = None,
    ) -> Dict[str, Any]:
        fields = self.extract_fields_local(text)
        fields["fallback_reason"] = reason
        if detail:
            fields["fallback_detail"] = detail
        return fields

    def _parse_json_or_local(self, content: str, ocr_text: str, reason: str) -> Dict[str, Any]:
        parsed = parse_model_json_or_fallback(content, ocr_text, reason)
        validation_errors = validate_invoice_fields(parsed, self.invoice_schema_path)
        if validation_errors:
            return self._local_with_reason(
                ocr_text,
                f"{reason}_schema_invalid",
                "; ".join(validation_errors),
            )
        return parsed

    @staticmethod
    def is_policy_block(message: str) -> bool:
        markers = [
            "zscaler",
            "violates compliance category",
            "posting content to this website is not allowed",
            "<!doctype html",
            "dlp policy",
            "internet security by zscaler",
        ]
        lowered = message.lower()
        return any(marker in lowered for marker in markers)

    def _usage_details_from_response(
        self,
        model_name: str,
        response: Any | None,
    ) -> tuple[dict[str, int] | None, float | None, str]:
        usage = normalize_usage(getattr(response, "usage", None) if response is not None else None)
        cost_usd = calculate_usage_cost(model_name, usage, self.telemetry.pricing)
        return usage, cost_usd, "provider" if usage is not None else "unavailable"

    def extract_fields_gpt4omini(
        self,
        ocr_text: str,
    ) -> Tuple[Dict[str, Any], dict[str, int] | None, float | None, str]:
        if self.openai_client is None:
            return self._local_with_reason(ocr_text, "gpt_unavailable_local"), None, None, "unavailable"

        try:
            response = self.openai_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "You are a precise invoice field extraction engine."},
                    {"role": "user", "content": f"{self.invoice_prompt}\n\nOCR_TEXT:\n{ocr_text}"},
                ],
                temperature=0,
                max_tokens=1000,
            )
            content = (response.choices[0].message.content or "").strip()
            usage, cost_usd, usage_source = self._usage_details_from_response(self.model_name, response)
            if self.is_policy_block(content):
                return (
                    self._local_with_reason(ocr_text, "gpt_policy_block_local"),
                    usage,
                    cost_usd,
                    usage_source,
                )
            return (
                self._parse_json_or_local(content, ocr_text, "gpt_parse_fallback_local"),
                usage,
                cost_usd,
                usage_source,
            )
        except Exception as exc:
            if self.is_policy_block(str(exc)):
                return self._local_with_reason(ocr_text, "gpt_policy_block_local"), None, None, "unavailable"
            return self._local_with_reason(ocr_text, "gpt_error_local", str(exc)), None, None, "unavailable"

    def _ensure_qwen_loaded(self) -> bool:
        if self._qwen_model is not None and self._qwen_processor is not None:
            return True
        if self._qwen_load_error is not None:
            return False
        if not self.qwen_model_dir.exists():
            self._qwen_load_error = f"Qwen model directory not found: {self.qwen_model_dir}"
            return False

        try:
            import torch
            from transformers import AutoProcessor, Qwen2VLForConditionalGeneration

            if torch.cuda.is_available():
                device = "cuda"
            elif hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
                device = "mps"
            else:
                device = "cpu"

            self._qwen_processor = AutoProcessor.from_pretrained(
                str(self.qwen_model_dir),
                local_files_only=True,
                trust_remote_code=True,
            )
            self._qwen_model = Qwen2VLForConditionalGeneration.from_pretrained(
                str(self.qwen_model_dir),
                local_files_only=True,
                trust_remote_code=True,
                torch_dtype="auto",
            )
            self._qwen_model.to(device)
            self._qwen_model.eval()
            self._qwen_device = device
            return True
        except Exception as exc:
            self._qwen_load_error = str(exc)
            return False

    def extract_fields_qwen(self, ocr_text: str) -> Dict[str, Any]:
        if not self._ensure_qwen_loaded():
            return self._local_with_reason(ocr_text, "qwen_unavailable_local", self._qwen_load_error)

        try:
            import torch

            messages = [
                {"role": "system", "content": "You are a precise invoice field extraction engine."},
                {"role": "user", "content": f"{self.invoice_prompt}\n\nOCR_TEXT:\n{ocr_text}"},
            ]
            chat_text = self._qwen_processor.apply_chat_template(
                messages,
                tokenize=False,
                add_generation_prompt=True,
            )
            inputs = self._qwen_processor(text=[chat_text], padding=True, return_tensors="pt")
            inputs = {
                key: (value.to(self._qwen_device) if hasattr(value, "to") else value)
                for key, value in inputs.items()
            }

            with torch.inference_mode():
                generated_ids = self._qwen_model.generate(
                    **inputs,
                    max_new_tokens=1000,
                    do_sample=False,
                    temperature=0.0,
                )

            trimmed_ids = [
                output_ids[len(input_ids) :]
                for input_ids, output_ids in zip(inputs["input_ids"], generated_ids)
            ]
            content = self._qwen_processor.batch_decode(
                trimmed_ids,
                skip_special_tokens=True,
                clean_up_tokenization_spaces=False,
            )[0].strip()
            return self._parse_json_or_local(content, ocr_text, "qwen_parse_fallback_local")
        except Exception as exc:
            return self._local_with_reason(ocr_text, "qwen_error_local", str(exc))

    def _build_llmops_metadata(
        self,
        provider: str,
        model_name: str,
        fields: Dict[str, Any],
        started_at: float,
        usage: dict[str, int] | None,
        cost_usd: float | None,
        usage_source: str,
        surface: str,
    ) -> Dict[str, Any]:
        validation_errors = validate_invoice_fields(fields, self.invoice_schema_path)
        return {
            "provider": provider,
            "model": model_name,
            "prompt_version": self.prompt_version,
            "schema_version": self.schema_version,
            "validation_status": "valid" if not validation_errors else "invalid",
            "validation_errors": validation_errors,
            "fallback_reason": fields.get("fallback_reason"),
            "latency_ms": round((time.perf_counter() - started_at) * 1000, 2),
            "prompt_tokens": usage["prompt_tokens"] if usage is not None else None,
            "completion_tokens": usage["completion_tokens"] if usage is not None else None,
            "total_tokens": usage["total_tokens"] if usage is not None else None,
            "cost_usd": cost_usd,
            "usage_source": usage_source,
            "surface": surface,
        }

    @staticmethod
    def _status_from_metadata(llmops_metadata: Dict[str, Any]) -> str:
        if llmops_metadata.get("fallback_reason"):
            return "fallback"
        if llmops_metadata.get("validation_status") != "valid":
            return "invalid"
        return "valid"

    def extract_fields_with_mode(self, ocr_text: str) -> Tuple[Dict[str, Any], str, Dict[str, Any]]:
        started_at = time.perf_counter()

        if self.field_extractor_mode == "gpt":
            fields, usage, cost_usd, usage_source = self.extract_fields_gpt4omini(ocr_text)
            extraction_mode = "gpt-4o-mini"
            if fields.get("fallback_reason"):
                extraction_mode = "local_fallback_after_gpt"
            return (
                fields,
                extraction_mode,
                self._build_llmops_metadata(
                    "openai",
                    extraction_mode,
                    fields,
                    started_at,
                    usage,
                    cost_usd,
                    usage_source,
                    "streamlit_app",
                ),
            )

        if self.field_extractor_mode == "qwen":
            fields = self.extract_fields_qwen(ocr_text)
            extraction_mode = "qwen-vl-local"
            if fields.get("fallback_reason"):
                extraction_mode = "local_fallback_after_qwen"
            return (
                fields,
                extraction_mode,
                self._build_llmops_metadata(
                    "qwen",
                    extraction_mode,
                    fields,
                    started_at,
                    None,
                    None,
                    "unavailable",
                    "streamlit_app",
                ),
            )

        if self.openai_client is not None:
            gpt_fields, usage, cost_usd, usage_source = self.extract_fields_gpt4omini(ocr_text)
            if not gpt_fields.get("fallback_reason"):
                return (
                    gpt_fields,
                    "gpt-4o-mini",
                    self._build_llmops_metadata(
                        "openai",
                        "gpt-4o-mini",
                        gpt_fields,
                        started_at,
                        usage,
                        cost_usd,
                        usage_source,
                        "streamlit_app",
                    ),
                )

        qwen_fields = self.extract_fields_qwen(ocr_text)
        if not qwen_fields.get("fallback_reason"):
            return (
                qwen_fields,
                "qwen-vl-local",
                self._build_llmops_metadata(
                    "qwen",
                    "qwen-vl-local",
                    qwen_fields,
                    started_at,
                    None,
                    None,
                    "unavailable",
                    "streamlit_app",
                ),
            )

        return (
            qwen_fields,
            "local_fallback",
            self._build_llmops_metadata(
                "local",
                "local_fallback",
                qwen_fields,
                started_at,
                None,
                None,
                "unavailable",
                "streamlit_app",
            ),
        )

    def _record_trace(
        self,
        source_name: str,
        document_type: str,
        ocr_text: str,
        fields: Dict[str, Any],
        extraction_mode: str,
        llmops_metadata: Dict[str, Any],
    ) -> None:
        write_trace_record(
            trace_path=self.trace_path,
            include_text=self.trace_include_text,
            record={
                "timestamp_utc": datetime.now(timezone.utc).isoformat(),
                "source_name": source_name,
                "document_type": document_type,
                "ocr_text": ocr_text,
                "fields": fields,
                "extraction_mode": extraction_mode,
                "llmops": llmops_metadata,
            },
        )

    def _record_request_metrics(
        self,
        document_type: str,
        llmops_metadata: Dict[str, Any],
    ) -> None:
        self.telemetry.record_request(
            surface=str(llmops_metadata.get("surface", "streamlit_app")),
            provider=str(llmops_metadata.get("provider", "unknown")),
            model=str(llmops_metadata.get("model", "unknown")),
            status=self._status_from_metadata(llmops_metadata),
            document_type=document_type,
            latency_ms=float(llmops_metadata.get("latency_ms", 0.0)),
            prompt_tokens=llmops_metadata.get("prompt_tokens"),
            completion_tokens=llmops_metadata.get("completion_tokens"),
            total_tokens=llmops_metadata.get("total_tokens"),
            cost_usd=llmops_metadata.get("cost_usd"),
        )

    def _missing_ocr_result(self) -> Dict[str, Any]:
        return {
            "status": "error",
            "error_code": "ocr_models_missing",
            "error": self.ocr_unavailable_reason or "Required OCR models are unavailable",
        }

    def ocr_jpg_upload(self, uploaded_file) -> Dict[str, Any]:
        if not self.ocr_available:
            return self._missing_ocr_result()

        file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
        bgr = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        if bgr is None:
            return {"status": "error", "error": "Image decode failed"}

        rgb = cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)
        metrics = self.quality_metrics(rgb)
        text, confidence, detections, layout_meta = self.ocr_image_layout_aware(rgb)
        fields, extraction_mode, llmops_metadata = self.extract_fields_with_mode(text)
        self._record_trace(uploaded_file.name, "jpg", text, fields, extraction_mode, llmops_metadata)
        self._record_request_metrics("jpg", llmops_metadata)

        return {
            "status": "success",
            "type": "jpg",
            "avg_confidence": confidence,
            "detections": detections,
            "quality": metrics,
            "text": text,
            "fields": fields,
            "extraction_mode": extraction_mode,
            "layout": layout_meta,
            "llmops": llmops_metadata,
        }

    def ocr_pdf_upload(self, uploaded_file) -> Dict[str, Any]:
        if not self.ocr_available:
            return self._missing_ocr_result()

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = Path(tmp.name)

        try:
            doc = pdfium.PdfDocument(str(tmp_path))
            page_outputs: List[Dict[str, Any]] = []
            all_text: List[str] = []
            all_confidences: List[float] = []

            for index in range(len(doc)):
                arr = np.array(doc[index].render(scale=2.0).to_pil().convert("RGB"))
                text, confidence, detections, layout_meta = self.ocr_image_layout_aware(arr)
                page_outputs.append(
                    {
                        "page": index + 1,
                        "avg_confidence": confidence,
                        "detections": detections,
                        "layout_status": layout_meta.get("layout_status"),
                        "layout_regions": layout_meta.get("layout_regions", 0),
                        "layout_crop_chunks": layout_meta.get("layout_crop_chunks", 0),
                    }
                )
                all_text.append(f"=== PAGE {index + 1} ===\n{text}")
                if detections > 0:
                    all_confidences.append(confidence)

            merged = "\n\n".join(all_text)
            fields, extraction_mode, llmops_metadata = self.extract_fields_with_mode(merged)
            self._record_trace(uploaded_file.name, "pdf", merged, fields, extraction_mode, llmops_metadata)
            self._record_request_metrics("pdf", llmops_metadata)

            return {
                "status": "success",
                "type": "pdf",
                "pages": len(doc),
                "avg_confidence": (
                    float(sum(all_confidences) / len(all_confidences)) if all_confidences else 0.0
                ),
                "page_stats": page_outputs,
                "layout_summary": {
                    "regions_total": int(sum(page.get("layout_regions", 0) for page in page_outputs)),
                    "crop_chunks_total": int(sum(page.get("layout_crop_chunks", 0) for page in page_outputs)),
                },
                "text": merged,
                "fields": fields,
                "extraction_mode": extraction_mode,
                "llmops": llmops_metadata,
            }
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except OSError:
                pass

    def ocr_docx_upload(self, uploaded_file) -> Dict[str, Any]:
        if not self.ocr_available:
            return self._missing_ocr_result()

        import io
        import zipfile

        import docx

        raw = uploaded_file.read()
        doc_obj = docx.Document(io.BytesIO(raw))
        native_parts: List[str] = []
        for para in doc_obj.paragraphs:
            text = para.text.strip()
            if text:
                native_parts.append(text)
        for table in doc_obj.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    native_parts.append(" | ".join(row_cells))
        native_text = "\n".join(native_parts)

        image_texts: List[str] = []
        image_stats: List[Dict[str, Any]] = []
        supported_exts = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"}

        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            media_files = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/") and Path(name).suffix.lower() in supported_exts
            ]
            for index, media_name in enumerate(media_files):
                img_bytes = archive.read(media_name)
                arr = cv2.imdecode(np.frombuffer(img_bytes, np.uint8), cv2.IMREAD_COLOR)
                if arr is None:
                    continue
                rgb = cv2.cvtColor(arr, cv2.COLOR_BGR2RGB)
                text, confidence, detections, layout_meta = self.ocr_image_layout_aware(rgb)
                image_texts.append(
                    f"=== EMBEDDED IMAGE {index + 1} ({Path(media_name).name}) ===\n{text}"
                )
                image_stats.append(
                    {
                        "image": Path(media_name).name,
                        "avg_confidence": confidence,
                        "detections": detections,
                        "layout_status": layout_meta.get("layout_status"),
                        "layout_regions": layout_meta.get("layout_regions", 0),
                        "layout_crop_chunks": layout_meta.get("layout_crop_chunks", 0),
                    }
                )

        sections: List[str] = []
        if native_text.strip():
            sections.append(f"=== DOCUMENT TEXT ===\n{native_text}")
        sections.extend(image_texts)
        merged = "\n\n".join(sections) if sections else native_text

        fields, extraction_mode, llmops_metadata = self.extract_fields_with_mode(merged)
        self._record_trace(uploaded_file.name, "docx", merged, fields, extraction_mode, llmops_metadata)
        self._record_request_metrics("docx", llmops_metadata)

        return {
            "status": "success",
            "type": "docx",
            "native_text_lines": len(native_parts),
            "embedded_images": len(image_texts),
            "image_stats": image_stats,
            "text": merged,
            "fields": fields,
            "extraction_mode": extraction_mode,
            "llmops": llmops_metadata,
        }

    def process_upload(self, uploaded_file) -> Dict[str, Any]:
        suffix = Path(uploaded_file.name).suffix.lower()
        if suffix == ".pdf":
            return self.ocr_pdf_upload(uploaded_file)
        if suffix == ".docx":
            return self.ocr_docx_upload(uploaded_file)
        return self.ocr_jpg_upload(uploaded_file)
